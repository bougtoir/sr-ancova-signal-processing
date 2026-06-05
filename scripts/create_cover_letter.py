#!/usr/bin/env python3
"""Generate cover letter for PRE submission."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent.parent / 'pre_submission'


def build_cover_letter():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run('[Date]')

    doc.add_paragraph()

    # Addressee
    p = doc.add_paragraph()
    p.add_run('Editor\nPhysical Review E\nAmerican Physical Society')

    doc.add_paragraph()

    # Subject
    p = doc.add_paragraph()
    run = p.add_run(
        'Re: Submission of manuscript "Covariate-adjusted stochastic resonance '
        'in threshold-based event detectors"'
    )
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Dear Editor,')

    doc.add_paragraph()

    doc.add_paragraph(
        'We submit for your consideration the enclosed manuscript entitled '
        '"Covariate-adjusted stochastic resonance in threshold-based event detectors" '
        'as a Regular Article in Physical Review E.'
    )

    doc.add_paragraph(
        'This work bridges two previously disconnected research traditions: stochastic '
        'resonance (SR) theory, which demonstrates that noise can enhance signal detection '
        'in nonlinear threshold systems, and covariate adjustment methods (analogous to '
        'ANCOVA in statistics), which provide systematic tools for noise reduction when '
        'noise depends on observable parameters. We derive an analytical expression for '
        'the optimal noise model accuracy — the degree to which noise should be modeled '
        'and removed — showing that the optimal strategy is to reduce noise to the SR '
        'optimum, not to zero. This result, ρ* = √(1 − θ²/σ²), provides a constructive '
        'criterion that complements existing existential results such as the '
        'forbidden-interval theorem.'
    )

    doc.add_paragraph(
        'The practical relevance of the framework is demonstrated through application to '
        'dynamic vision sensor (DVS) data from astronomical observations, where a '
        'physics-informed noise model achieves ROC-AUC = 0.866 for noise classification '
        'with 93.9% signal preservation — performance consistent with the theoretical '
        'predictions.'
    )

    doc.add_paragraph(
        'We believe this work is well suited to Physical Review E given the journal\'s '
        'strong tradition in stochastic resonance, nonlinear dynamics, and statistical '
        'physics. The manuscript connects fundamental SR theory with practical signal '
        'processing methodology, which we expect to be of interest to readers across '
        'the fields of nonlinear physics, sensor engineering, and computational neuroscience.'
    )

    doc.add_paragraph(
        'This manuscript has not been published or submitted for publication elsewhere. '
        'All authors have approved the manuscript and agree to its submission.'
    )

    doc.add_paragraph()

    doc.add_paragraph('Sincerely,')
    doc.add_paragraph('[Author names]')

    out_path = OUT_DIR / 'cover_letter_pre.docx'
    doc.save(str(out_path))
    print(f"Cover letter saved: {out_path}")


if __name__ == '__main__':
    build_cover_letter()
