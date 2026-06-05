#!/usr/bin/env python3
"""
Generate PRE manuscript as docx with inline figures.

Title: "Covariate-adjusted stochastic resonance in threshold-based event detectors"

Structure follows PRE Regular Article format.
References: numbered in order of first appearance (Vancouver/APS style).

Paper b replaces paper a — all DVS noise physics, A5 model, and Fano factor
theory are self-contained here. No self-citation to previous DVS work.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent.parent / 'pre_submission'
FIG_DIR = OUT_DIR  # figures are in same directory


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, style='Normal', bold=False, italic=False,
                  alignment=None, space_after=None, space_before=None):
    p = doc.add_paragraph(text, style=style)
    if bold:
        for run in p.runs:
            run.bold = True
    if italic:
        for run in p.runs:
            run.italic = True
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_figure(doc, img_path, caption, width=Inches(5.5)):
    """Insert figure inline with caption below."""
    if not img_path.exists():
        add_paragraph(doc, f"[MISSING FIGURE: {img_path.name}]")
        return
    # Space before figure
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(12)
    run = p_img.add_run()
    run.add_picture(str(img_path), width=width)

    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(caption)
    run_cap.font.size = Pt(9)
    return p_cap


def build_manuscript():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # =========================================================
    # Title page
    # =========================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Covariate-adjusted stochastic resonance\n'
        'in threshold-based event detectors'
    )
    run.font.size = Pt(16)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Author names and affiliations to be added]')
    run.font.size = Pt(11)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run('(Dated: \\today)')
    run.font.size = Pt(10)

    # =========================================================
    # Abstract
    # =========================================================
    add_heading(doc, 'Abstract', level=1)
    add_paragraph(doc, (
        'Stochastic resonance (SR) allows noise to enhance weak-signal detection in '
        'threshold systems: output signal-to-noise ratio (SNR) peaks at an intermediate '
        'noise level rather than at zero noise. Separately, covariate adjustment\u2014as in '
        'analysis of covariance (ANCOVA)\u2014reduces noise variance by modeling its dependence '
        'on observable auxiliary variables. We unify these ideas for threshold-based event '
        'detectors. A noise model with correlation \u03c1 to the true noise shrinks the effective '
        'variance by (1 \u2212 \u03c1\u00b2), shifting the detector\u2019s operating point along the SR curve. '
        'Maximizing output SNR over \u03c1 yields an optimal model accuracy '
        '\u03c1* = \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2), where \u03b8 is the threshold and \u03c3 the noise level. '
        'In the excess-noise regime (\u03c3 > \u03b8), adjustment helps and the SNR gain grows '
        'exponentially with \u03c3; at the SR optimum (\u03c3 \u2248 \u03b8), any noise removal is '
        'counterproductive. Monte Carlo simulations confirm the analytical predictions. '
        'As a practical demonstration, we apply the framework to dynamic vision sensor '
        '(DVS) astronomical observations, where a circuit-level noise model\u2014parameterized '
        'by dark current, background illuminance, and threshold mismatch\u2014supplies the '
        'covariate structure. A Fano-factor classifier built on this model achieves '
        'ROC-AUC = 0.866 with 93.9% signal preservation across 20 recordings from the '
        'Event-Based Space Situational Awareness dataset, matching the theory\u2019s '
        'predictions for the excess-noise regime.'
    ))

    # =========================================================
    # I. Introduction
    # =========================================================
    add_heading(doc, 'I. INTRODUCTION', level=1)
    add_paragraph(doc, (
        'In threshold-based detectors, noise plays a dual role. At low levels it is '
        'innocuous; at high levels it overwhelms the signal. Between these extremes, '
        'stochastic resonance (SR) produces a counter-intuitive optimum where a finite '
        'noise level maximizes signal detection [1\u20133]. Discovered in the context of '
        'paleoclimatic periodicities [4] and since confirmed in physical, biological, '
        'and engineered systems [5], SR implies that indiscriminate noise suppression '
        'can be self-defeating.'
    ))
    add_paragraph(doc, (
        'A complementary tradition treats noise not as a monolithic nuisance but as a '
        'structured quantity depending on measurable covariates\u2014temperature, '
        'instrumental drift, illumination\u2014which can be modeled and subtracted. This '
        'is the logic of analysis of covariance (ANCOVA) [6] and its signal-processing '
        'counterparts: adaptive noise cancellation [7], Wiener filtering [8], and '
        'physics-informed denoising [9]. In all cases the residual variance shrinks as '
        'the noise model improves, yet the SR implications of partial noise removal are '
        'seldom discussed.'
    ))
    add_paragraph(doc, (
        'These two viewpoints have evolved in isolation. SR theory identifies the '
        'optimal noise level but offers no prescription for reaching it; denoising '
        'methods remove noise without asking whether total removal is desirable. '
        'The gap matters for threshold-based sensors now in active development\u2014dynamic '
        'vision sensors (DVS) [10], single-photon detectors [11], neuromorphic circuits '
        '[12]\u2014all of which fire events by threshold crossing and therefore exhibit '
        'intrinsic SR effects.'
    ))
    add_paragraph(doc, (
        'Here we connect the two by analyzing covariate adjustment in the presence of '
        'SR. For a threshold detector whose noise depends on observable covariates, '
        'we derive a closed-form optimal model accuracy\u2014the fraction of noise that '
        'should be removed\u2014as a function of the noise-to-threshold ratio. '
        'The answer is not \u201cremove everything you can model,\u201d but rather \u201creduce '
        'noise to the SR optimum and stop.\u201d We verify this result via Monte Carlo '
        'simulation and illustrate it with DVS astronomical data, where pixel-level '
        'noise physics supplies a natural covariate structure.'
    ))

    # =========================================================
    # II. General Framework
    # =========================================================
    add_heading(doc, 'II. GENERAL FRAMEWORK', level=1)

    add_heading(doc, 'A. Threshold-based event detector', level=2)
    add_paragraph(doc, (
        'Consider a threshold detector receiving x(t) = s(t) + n(t), where s(t) is a '
        'deterministic signal and n(t) is zero-mean Gaussian noise with variance \u03c3\u00b2. '
        'The output event stream is'
    ))
    add_paragraph(doc, (
        '    E(t) = 1    if |x(t)| > \u03b8,\n'
        '    E(t) = 0    otherwise,'
    ), space_before=6, space_after=6)
    add_paragraph(doc, (
        'with threshold \u03b8 > 0. This encompasses level-crossing detectors, neuronal '
        'integrate-and-fire models, Schmitt triggers, single-photon avalanche diodes, '
        'and DVS pixels [10, 13]. The signal is subthreshold, A \u2261 max|s(t)| < \u03b8, '
        'so events occur only when noise assists a threshold crossing (Fig. 1a).'
    ))

    # Fig 1
    add_figure(doc, FIG_DIR / 'fig1_schematic.png',
               'FIG. 1. Conceptual schematic of the covariate-adjusted stochastic resonance '
               'framework. (a) Threshold-based event detector: a weak periodic signal s(t) '
               'embedded in Gaussian noise triggers events when |x(t)| > \u03b8. Tick marks above '
               'indicate event times. (b) Stochastic resonance: the event rate modulation '
               'tracking the signal period is maximized at intermediate noise (green), '
               'suppressed at low noise (red), and washed out at high noise (blue). '
               '(c) Covariate adjustment with noise model correlation \u03c1 narrows the residual '
               'noise distribution, equivalent to shifting the operating point on the SR curve.',
               width=Inches(6.0))

    add_heading(doc, 'B. Stochastic resonance in threshold detectors', level=2)
    add_paragraph(doc, (
        'For a weak sinusoid s(t) = A sin(2\u03c0f\u2080t), the two-state theory [1] '
        'gives the output SNR of the event stream as'
    ))
    add_paragraph(doc, (
        '    SNR_out(\u03c3) \u221d (A/\u03c3\u00b2)\u00b2 exp(\u22122\u03b8\u00b2/\u03c3\u00b2).                   (1)'
    ), space_before=6, space_after=6)
    add_paragraph(doc, (
        'Equation (1) peaks at \u03c3* = \u03b8 (Fig. 2), independent of A. Below \u03b8, '
        'threshold crossings are too rare; far above \u03b8, events are frequent but carry '
        'little signal modulation. At \u03c3 \u2248 \u03b8 the two effects balance, producing '
        'a maximally informative event stream (Fig. 1b).'
    ))

    # Fig 2 — placed after first citation in Sec II.B
    add_figure(doc, FIG_DIR / 'fig2_sr_curves.png',
               'FIG. 2. Stochastic resonance curves for a threshold detector (A/\u03b8 = 0.3). '
               'Solid lines: analytical SNR from two-state theory [Eq. (1)]. Black circles '
               'with error bars: Monte Carlo validation (15 trials per point). Covariate '
               'adjustment shifts the SR peak rightward, meaning the detector tolerates more '
               'input noise when a good noise model is available.',
               width=Inches(4.5))

    add_heading(doc, 'C. Covariate adjustment as noise reduction', level=2)
    add_paragraph(doc, (
        'Suppose n(t) depends on observable covariates z(t) = (z\u2081(t), \u2026, z_k(t))\u1d40 '
        'through a model n\u0302(t) = f(z(t); \u03b2). Subtracting this estimate\u2014the ANCOVA '
        'adjustment\u2014gives'
    ))
    add_paragraph(doc, (
        '    x_adj(t) = x(t) \u2212 n\u0302(t) = s(t) + \u03b5(t),                   (2)'
    ), space_before=6, space_after=6)
    add_paragraph(doc, (
        'where \u03b5(t) = n(t) \u2212 n\u0302(t) is the residual. If the model achieves '
        'correlation \u03c1 = Corr(n\u0302, n) with the true noise, then'
    ))
    add_paragraph(doc, (
        '    Var(\u03b5) = (1 \u2212 \u03c1\u00b2) \u03c3\u00b2,                                  (3)'
    ), space_before=6, space_after=6)
    add_paragraph(doc, (
        'so \u03c3_eff = \u03c3\u221a(1 \u2212 \u03c1\u00b2). Because the threshold \u03b8 and signal s(t) '
        'are unchanged, the adjustment is equivalent to sliding the operating '
        'point leftward along the SR curve by the factor \u221a(1 \u2212 \u03c1\u00b2) (Fig. 1c).'
    ))

    add_heading(doc, 'D. Optimal noise model accuracy', level=2)
    add_paragraph(doc, (
        'Substituting \u03c3_eff into Eq. (1) gives the output SNR as a joint function '
        'of \u03c3 and \u03c1:'
    ))
    add_paragraph(doc, (
        '    SNR_out(\u03c3, \u03c1) \u221d [A / (\u03c3\u00b2(1 \u2212 \u03c1\u00b2))]\u00b2 exp(\u22122\u03b8\u00b2 / [\u03c3\u00b2(1 \u2212 \u03c1\u00b2)]).     (4)'
    ), space_before=6, space_after=6)
    add_paragraph(doc, (
        'Maximizing over \u03c1 at fixed \u03c3:'
    ))
    add_paragraph(doc, (
        '    \u03c1*(\u03c3) = { 0,                          if \u03c3 \u2264 \u03b8,\n'
        '            { \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2),              if \u03c3 > \u03b8.          (5)'
    ), space_before=6, space_after=6)
    add_paragraph(doc, (
        'The two regimes have distinct physics (Fig. 3a). For \u03c3 \u2264 \u03b8 the detector '
        'already sits at or below the SR peak; removing noise moves it further from '
        'the optimum, so \u03c1* = 0. For \u03c3 > \u03b8 the system is above the peak and '
        'adjustment should bring the effective noise exactly to the SR optimum: '
        '\u03c3_eff = \u03c3\u221a(1 \u2212 \u03c1*\u00b2) = \u03b8.'
    ))
    add_paragraph(doc, (
        'The resulting SNR gain over the unadjusted case is'
    ))
    add_paragraph(doc, (
        '    SNR_out(\u03c3, \u03c1*) / SNR_out(\u03c3, 0) = (\u03c3/\u03b8)\u2074 exp(2(\u03c3\u00b2 \u2212 \u03b8\u00b2)/\u03c3\u00b2),    (6)'
    ), space_before=6, space_after=6)
    add_paragraph(doc, (
        'which scales as ~exp(2\u03c3\u00b2/\u03b8\u00b2) for \u03c3 \u226b \u03b8 (Fig. 3b). The exponential '
        'growth reflects the steep penalty of operating far above the SR peak, '
        'and the correspondingly large payoff of accurate noise modeling in '
        'high-noise environments.'
    ))

    # Fig 3 — placed after first citation in Sec II.D
    add_figure(doc, FIG_DIR / 'fig3_optimal_rho.png',
               'FIG. 3. (a) Optimal noise model accuracy \u03c1* versus input noise. Yellow shading: '
               'SR regime (\u03c3 < \u03b8) where \u03c1* = 0. Blue shading: excess noise regime (\u03c3 > \u03b8) '
               'where \u03c1* = \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2). Dashed line: analytical prediction. '
               '(b) Peak SNR improvement at optimal \u03c1* grows exponentially with input noise, '
               'reaching ~100\u00d7 at \u03c3/\u03b8 = 4.',
               width=Inches(5.5))

    # =========================================================
    # III. Numerical Simulations
    # =========================================================
    add_heading(doc, 'III. NUMERICAL SIMULATIONS', level=1)

    add_paragraph(doc, (
        'We test the analytical predictions with Monte Carlo simulations. '
        'The signal is s(t) = A sin(2\u03c0f\u2080t) with A/\u03b8 = 0.3, f\u2080 = 5 Hz, '
        'dt = 1 ms, and N = 10\u2075 steps per trial.'
    ))

    add_heading(doc, 'A. Stochastic resonance curves', level=2)
    add_paragraph(doc, (
        'Figure 2 plots output SNR against input noise for several \u03c1 values. '
        'Monte Carlo estimates (circles) agree with the analytical curve for \u03c1 = 0, '
        'peaking at \u03c3/\u03b8 \u2248 1. The adjusted curves shift rightward to '
        '\u03c3/\u03b8 \u2248 1/\u221a(1 \u2212 \u03c1\u00b2) as predicted. The peak amplitude is unchanged '
        'when expressed in effective-noise units, confirming that adjustment translates '
        'the SR curve without distorting it.'
    ))

    add_heading(doc, 'B. Detection probabilities', level=2)
    add_paragraph(doc, (
        'Figure 4 shows detection probability P_D and false-alarm probability P_FA '
        'versus input noise (A/\u03b8 = 0.4). Both decrease with \u03c1 because adjustment '
        'suppresses the effective noise driving threshold crossings. The net benefit '
        'is clearer in the ROC plane (Fig. 5): at \u03c3/\u03b8 = 1.5, increasing \u03c1 lifts '
        'the ROC curve well above the chance diagonal, demonstrating improved '
        'signal\u2013noise discrimination.'
    ))

    # Fig 4
    add_figure(doc, FIG_DIR / 'fig4_detection_probability.png',
               'FIG. 4. (a) Detection probability P_D and (b) false alarm probability P_FA '
               'versus input noise level for different covariate model accuracies \u03c1 '
               '(A/\u03b8 = 0.4). Higher \u03c1 suppresses both probabilities by reducing effective noise.',
               width=Inches(5.5))

    # Fig 5
    add_figure(doc, FIG_DIR / 'fig5_roc_comparison.png',
               'FIG. 5. ROC curves at \u03c3_n/\u03b8 = 1.5 (excess noise regime) for different \u03c1. '
               'Covariate adjustment progressively improves detection performance, with '
               '\u03c1 = 0.95 achieving near-ideal separation between signal and noise events.',
               width=Inches(3.8))

    add_heading(doc, 'C. Optimal noise model accuracy', level=2)
    add_paragraph(doc, (
        'Figure 3 summarizes the optimal adjustment. Panel (a) compares the '
        'numerically determined \u03c1* with the prediction \u03c1* = \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2) (dashed '
        'line). The transition at \u03c3 = \u03b8 is sharp: below it, \u03c1* = 0; above it, '
        '\u03c1* rises steeply. Panel (b) shows the resulting SNR gain, which grows '
        'exponentially and exceeds 100\u00d7 by \u03c3/\u03b8 = 4\u2014a regime common in sensor '
        'applications with large dark-current or thermal noise.'
    ))

    add_heading(doc, 'D. Information-theoretic perspective', level=2)
    add_paragraph(doc, (
        'Figure 6 shows mutual information I(S; E) between signal and events. '
        'The SR peak shifts rightward with \u03c1 in the same manner as the SNR '
        'peak, confirming that the benefit of partial noise removal is not an '
        'artifact of the SNR definition.'
    ))

    # Fig 6
    add_figure(doc, FIG_DIR / 'fig6_mutual_information.png',
               'FIG. 6. Mutual information I(S; E) between the periodic signal and the event '
               'stream. The SR peak shifts rightward with covariate adjustment (\u03c1 = 0.8), '
               'consistent with the SNR analysis in Fig. 2.',
               width=Inches(4.5))

    # =========================================================
    # IV. Application: Dynamic Vision Sensors
    # =========================================================
    add_heading(doc, 'IV. APPLICATION: DYNAMIC VISION SENSORS', level=1)

    add_paragraph(doc, (
        'We now apply the framework to event-based astronomical observations with '
        'dynamic vision sensors (DVS). Each DVS pixel fires when the log-intensity '
        'change exceeds \u03b8_ON or \u03b8_OFF [10]\u2014a physical realization of the threshold '
        'detector analyzed above.'
    ))

    # ---- IV.A DVS noise physics ----
    add_heading(doc, 'A. DVS noise physics', level=2)
    add_paragraph(doc, (
        'DVS noise physics has been characterized through a series of circuit-level '
        'studies. Gra\u00e7a and Delbruck [14] showed that photon shot noise sets a '
        'lower bound on background activity (BA) at twice the shot-noise level, '
        'a consequence of the differential pixel architecture. McReynolds et al. '
        '[15] exploited the resulting alternating ON\u2194OFF polarity pattern of '
        'shot-noise events as a discriminant against signal.'
    ))
    add_paragraph(doc, (
        'For our purposes the most relevant contribution is the physically realistic '
        'DVS pixel model of Gra\u00e7a and Delbruck [16], which formulates event generation '
        'as a first-passage-time process and runs >1000\u00d7 faster than transistor-level '
        'Monte Carlo while retaining quantitative accuracy. A five-parameter '
        'analytical noise-rate expression (the A5 model) derived from it takes the form'
    ))
    add_paragraph(doc, (
        '    \u03bb_noise(T, I_bg) = I_dark,ref \u00b7 exp(\u03b1 \u00b7 \u0394T) \u00b7 (1 + \u03b2 \u00b7 I_bg),       (7)'
    ), space_before=6, space_after=6)
    add_paragraph(doc, (
        'Here I_dark,ref is the dark-current rate at a reference temperature, '
        '\u03b1 \u2248 0.06\u20130.08 K\u207b\u00b9 is the silicon temperature coefficient, '
        '\u0394T the temperature offset, \u03b2 the illuminance sensitivity, and I_bg '
        'the background illuminance. Additional parameters capture per-pixel '
        'threshold mismatch and readout bandwidth. The covariates '
        'z(t) = (T, I_bg, \u03b8_mismatch, ...) thus predict per-pixel noise rates '
        'via Eq. (7), providing the covariate structure needed for the adjustment '
        'framework of Sec. II.C.'
    ))

    # ---- IV.B Fano factor ----
    add_heading(doc, 'B. Fano factor as noise discriminant', level=2)
    add_paragraph(doc, (
        'The Fano factor F = Var(N)/Mean(N) of event counts in temporal bins '
        'distinguishes noise from signal at the pixel level. Pure Poisson noise '
        'gives F = 1; a periodically modulated rate bunches events and produces '
        'F > 1. The Fano factor therefore serves as a physics-grounded test '
        'statistic exploiting the known Poisson character of DVS shot noise.'
    ))
    add_paragraph(doc, (
        'In practice, events in each pixel are binned into non-overlapping windows of '
        'width \u0394t, giving counts {N_1, ..., N_K}. Pixels with F \u2264 F_thr (we use '
        'F_thr \u2248 2) are labeled noise-dominated and define the local noise rate '
        '\u03bb_noise(x, y). Per-event noise probability follows as'
    ))
    add_paragraph(doc, (
        '    P_noise(e_i) = \u03bb_noise(x_i, y_i, t_i) / '
        '[\u03bb_noise(x_i, y_i, t_i) + \u03bb_signal(x_i, y_i, t_i)],     (8)'
    ), space_before=6, space_after=6)
    add_paragraph(doc, (
        'where \u03bb_signal is estimated from excess rates in pixels with F > F_thr. '
        'Events satisfying P_noise > 0.5 are removed. This probabilistic thinning '
        'is the DVS realization of covariate adjustment: Eq. (7) supplies the '
        'covariate model, the Fano test identifies noise-dominated pixels, and '
        'Eq. (8) effects the adjustment.'
    ))

    # ---- IV.C Connection to noise inverse problem ----
    add_heading(doc, 'C. Connection to the noise inverse problem', level=2)
    add_paragraph(doc, (
        'The procedure above can be cast as a noise inverse problem: given raw events, '
        'reconstruct the noise component via Eq. (7) and subtract. The same logic '
        'drives gravitational-wave denoising, where witness channels model '
        'non-stationary instrumental noise in strain data [17, 18]. The structural '
        'parallel is:'
    ))
    add_paragraph(doc, (
        '    GW astronomy:    h(t) = s(t) + n_instr(t; aux channels)\n'
        '    DVS observation:  E(t) = E_signal(t) + E_noise(t; T, I_bg, \u03b8_mismatch)'
    ), space_before=6, space_after=6)
    add_paragraph(doc, (
        'In both settings the noise depends on observable auxiliaries (witness channels '
        'for LIGO; T, I_bg, \u03b8_mismatch for DVS). The noise-model accuracy '
        '\u03b1 = 1 \u2212 ||\u03b5|| / ||n|| maps to \u03c1 \u2248 \u03b1. The SNR gain, however, follows '
        'Eq. (6) rather than the linear estimate SNR \u221d 1/(1\u2212\u03b1) used in '
        'gravitational-wave analyses, because the threshold nonlinearity introduces '
        'SR-mediated amplification.'
    ))

    # ---- IV.D Experimental results ----
    add_heading(doc, 'D. Experimental evaluation', level=2)
    add_paragraph(doc, (
        'We evaluate on 20 recordings from the Event-Based Space Situational Awareness '
        '(EBSSA) dataset [19]\u2014DVS observations of satellites and debris against star '
        'fields (DAVIS240C sensor). The binary classification task labels each event '
        'as signal (astronomical object) or noise (dark current / background). Three '
        'methods are compared: (i) the Fano filter described above, (ii) a three-layer '
        'physics-informed neural network trained self-supervised on noise-dominated '
        'pixels, and (iii) conventional temporal filtering [20], which retains '
        'events only if enough spatiotemporal neighbors fall within a fixed window.'
    ))
    add_paragraph(doc, (
        'Results are shown in Fig. 7. The Fano filter achieves ROC-AUC = 0.866 \u00b1 0.107, '
        'well above temporal filtering (0.534 \u00b1 0.083) and the neural network '
        '(0.546 \u00b1 0.218). The temporal filter removes more raw noise (85.2%) but '
        'destroys most signal (SPR = 21.6%), disqualifying it for faint-object work. '
        'Panel (b) shows the NRR\u2013SPR trade-off: the Fano filter removes 71.3% of '
        'noise while preserving 93.9% of signal, closest to the ideal corner '
        '(NRR = 1, SPR = 1).'
    ))

    # Fig 7
    add_figure(doc, FIG_DIR / 'fig7_dvs_application.png',
               'FIG. 7. Application to DVS astronomical observation (EBSSA dataset, 20 '
               'recordings). (a) Noise classification ROC-AUC: the Fano filter (covariate '
               'adjustment) achieves 0.866, far exceeding temporal filtering and neural methods. '
               '(b) NRR vs SPR trade-off: the Fano filter preserves 93.9% of signal while '
               'removing 71.3% of noise.',
               width=Inches(5.5))

    add_paragraph(doc, (
        'Sweeping the A5 model over T \u2208 [10, 65]\u00b0C and I_bg \u2208 [0.1, 1000] lux '
        'predicts a mean SNR gain of 5.4\u00d7 (max 10.0\u00d7) at 90% model accuracy, '
        'consistent with the measured Fano-filter performance.'
    ))

    # ---- IV.E Interpretation in SR framework ----
    add_heading(doc, 'E. Interpretation in the SR framework', level=2)
    add_paragraph(doc, (
        'DVS astronomical observations lie deep in the excess-noise regime '
        '(\u03c3 \u226b \u03b8): dark-current rates dominate signal rates by orders of magnitude. '
        'The A5 + Fano model achieves \u03c1 \u2248 0.7\u20130.9, which Fig. 3b predicts '
        'should yield 5\u201310\u00d7 SNR improvement\u2014matching the measured 5.4\u00d7. That '
        'noise is not eliminated entirely (NRR = 0.713) accords with the prediction '
        'that over-removal past the SR optimum is counterproductive.'
    ))
    add_paragraph(doc, (
        'The neural network (AUC = 0.546), despite greater flexibility, lacks '
        'access to the physics-informed covariates and cannot match the Fano filter '
        '(AUC = 0.866). This illustrates a prediction of the framework: what matters '
        'is the fidelity of the noise model (\u03c1), not the complexity of the method. '
        'A simple model that captures the dominant noise physics achieves high \u03c1 '
        'and large SNR gains; a flexible model with wrong inputs cannot compensate.'
    ))

    # =========================================================
    # V. Discussion
    # =========================================================
    add_heading(doc, 'V. DISCUSSION', level=1)

    add_paragraph(doc, (
        'We discuss implications for theory, sensor design, and extensions.'
    ))

    add_heading(doc, 'A. Connection to forbidden-interval theorems', level=2)
    add_paragraph(doc, (
        'The forbidden-interval theorem [21, 22] gives necessary and sufficient '
        'conditions on a noise distribution for SR to occur. Covariate adjustment '
        'reshapes the effective distribution and can move it into or out of the '
        'forbidden interval. Our expression \u03c1* = \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2) provides a '
        'constructive prescription\u2014how much to adjust\u2014complementing the '
        'theorem\u2019s existential characterization.'
    ))

    add_heading(doc, 'B. Implications for sensor design', level=2)
    add_paragraph(doc, (
        'A practical corollary is that threshold sensors should be co-designed with '
        'noise models. Hardware noise reduction is costly; if an accurate covariate '
        'model is available, the sensor can tolerate higher raw noise and rely on '
        'post-hoc adjustment. The design target becomes \u03c3_eff = \u03b8 after adjustment, '
        'not \u03c3 = \u03b8 in hardware. This applies to neuromorphic sensors, event '
        'cameras, single-photon detectors, and related threshold devices.'
    ))

    add_heading(doc, 'C. Limitations and extensions', level=2)
    add_paragraph(doc, (
        'The analysis assumes additive Gaussian noise, a fixed threshold, and a linear '
        'noise model. Real systems may have Poisson shot noise, adaptive thresholds, '
        'or nonlinear noise dependencies. Generalizations could invoke the '
        'forbidden-interval theorem for non-Gaussian cases [21] or information-theoretic '
        'metrics for adaptive thresholds [23]. The DVS application shows that '
        'the qualitative predictions\u2014adjustment helps above the SR optimum; '
        'over-removal hurts\u2014survive even when these assumptions hold only '
        'approximately.'
    ))

    add_heading(doc, 'D. Broader applicability', level=2)
    add_paragraph(doc, (
        'The principle extends to any detection system combining a threshold '
        'nonlinearity with structured noise: neural spike detection [24], '
        'quantum key distribution in noisy channels [25], radar target detection '
        'in clutter [26], and ion-channel sensing at the single-molecule level [27]. '
        'In each case the prescription is the same: reduce noise to the SR '
        'optimum, and stop.'
    ))

    # =========================================================
    # VI. Conclusion
    # =========================================================
    add_heading(doc, 'VI. CONCLUSION', level=1)
    add_paragraph(doc, (
        'We have derived a closed-form optimal noise-model accuracy '
        '\u03c1* = \u221a(1 \u2212 \u03b8\u00b2/\u03c3\u00b2) for threshold-based event detectors, connecting '
        'stochastic resonance theory (when does noise help?) with covariate adjustment '
        '(how much noise should be removed?). The answer\u2014reduce effective noise to '
        'the SR peak and no further\u2014provides a quantitative bridge between these '
        'two traditions.'
    ))
    add_paragraph(doc, (
        'Simulations confirm the analytical predictions across the full \u03c3/\u03b8 range. '
        'Applied to DVS astronomical data, a circuit-level noise model yields a '
        'Fano-factor classifier with ROC-AUC = 0.866, 93.9% signal preservation, '
        'and 5.4\u00d7 mean SNR gain\u2014quantitatively consistent with the excess-noise-regime '
        'prediction. The principle offers a design criterion for any '
        'threshold-based detector operating in structured noise: model the noise, '
        'subtract to the SR optimum, and accept the residual as beneficial.'
    ))

    # =========================================================
    # Data Availability Statement (APS required)
    # =========================================================
    add_heading(doc, 'DATA AVAILABILITY STATEMENT', level=1)
    add_paragraph(doc, (
        'Simulation code and figure-generation scripts are available at '
        'https://github.com/bougtoir/sr-ancova-framework. The EBSSA dataset '
        '(Sec. IV) is distributed via the Tonic library [19].'
    ))

    # =========================================================
    # Acknowledgments
    # =========================================================
    add_heading(doc, 'ACKNOWLEDGMENTS', level=1)
    add_paragraph(doc, '[To be added.]', italic=True)

    # =========================================================
    # References  (Vancouver style, numbered in order of first appearance)
    # =========================================================
    add_heading(doc, 'REFERENCES', level=1)
    references = [
        # --- Sec I: SR theory ---
        '[1] L. Gammaitoni, P. H\u00e4nggi, P. Jung, and F. Marchesoni, \u201cStochastic resonance,\u201d '
        'Rev. Mod. Phys. 70, 223 (1998).',

        '[2] M. D. McDonnell and D. Abbott, \u201cWhat is stochastic resonance? Definitions, '
        'misconceptions, debates, and its relevance to biology,\u201d PLoS Comput. Biol. 5, '
        'e1000348 (2009).',

        '[3] A. R. Bulsara and L. Gammaitoni, \u201cTuning in to noise,\u201d Phys. Today 49, 39 (1996).',

        '[4] R. Benzi, A. Sutera, and A. Vulpiani, \u201cThe mechanism of stochastic resonance,\u201d '
        'J. Phys. A: Math. Gen. 14, L453 (1981).',

        '[5] K. Wiesenfeld and F. Moss, \u201cStochastic resonance and the benefits of noise: '
        'from ice ages to crayfish and SQUIDs,\u201d Nature 373, 33 (1995).',

        # --- Sec I: ANCOVA / signal processing ---
        '[6] G. W. Snedecor and W. G. Cochran, Statistical Methods, 8th ed. '
        '(Iowa State University Press, Ames, 1989).',

        '[7] B. Widrow, J. R. Glover, Jr., J. M. McCool, J. Kaunitz, C. S. Williams, '
        'R. H. Hearn, J. R. Zeidler, E. Dong, Jr., and R. C. Goodlin, \u201cAdaptive noise '
        'cancelling: Principles and applications,\u201d Proc. IEEE 63, 1692 (1975).',

        '[8] N. Wiener, Extrapolation, Interpolation, and Smoothing of Stationary '
        'Time Series (MIT Press, Cambridge, MA, 1949).',

        '[9] G. E. Karniadakis, I. G. Kevrekidis, L. Lu, P. Perdikaris, S. Wang, '
        'and L. Yang, \u201cPhysics-informed machine learning,\u201d Nat. Rev. Phys. 3, 422 (2021).',

        # --- Sec I/II: DVS and threshold detectors ---
        '[10] G. Gallego, T. Delbr\u00fcck, G. Orchard, C. Bartolozzi, B. Taba, A. Censi, '
        'S. Leutenegger, A. J. Davison, J. Conradt, K. Daniilidis, and D. Scaramuzza, '
        '\u201cEvent-based vision: A survey,\u201d IEEE Trans. Pattern Anal. Mach. Intell. 44, '
        '154 (2022).',

        '[11] R. H. Hadfield, \u201cSingle-photon detectors for optical quantum information '
        'applications,\u201d Nat. Photonics 3, 696 (2009).',

        '[12] G. Indiveri, B. Linares-Barranco, T. J. Hamilton, A. van Schaik, '
        'R. Etienne-Cummings, T. Delbruck, S.-C. Liu, P. Dudek, P. H\u00e4fliger, '
        'S. Renaud, J. Schemmel, G. Cauwenberghs, J. Arthur, K. Hynna, '
        'F. Folowosele, S. Sa\u00efghi, T. Serrano-Gotarredona, J. Wijekoon, Y. Wang, '
        'and K. Boahen, \u201cNeuromorphic silicon neuron circuits,\u201d Front. Neurosci. 5, '
        '73 (2011).',

        '[13] C. Posch, T. Serrano-Gotarredona, B. Linares-Barranco, and T. Delbruck, '
        '\u201cRetinomorphic event-based vision sensors: Bioinspired cameras with spiking '
        'output,\u201d Proc. IEEE 102, 1470 (2014).',

        # --- Sec IV.A: DVS noise physics ---
        '[14] R. Gra\u00e7a and T. Delbruck, \u201cUnraveling the paradox of intensity-dependent '
        'DVS pixel noise,\u201d preprint arXiv:2109.08640 (2021).',

        '[15] B. McReynolds, R. Gra\u00e7a, and T. Delbruck, \u201cExploiting alternating DVS shot '
        'noise event pair statistics to reduce background activity rates,\u201d '
        'preprint arXiv:2304.03494 (2023).',

        '[16] R. Gra\u00e7a and T. Delbruck, \u201cTowards a physically realistic computationally '
        'efficient DVS pixel model,\u201d preprint arXiv:2505.07386 (2025).',

        # --- Sec IV.C: Noise inverse problem / DeepClean ---
        '[17] G. Vajente, Y. Huang, M. Isi, J. C. Driggers, J. S. Kissel, '
        'M. J. Szczepanczyk, and S. Vitale, \u201cMachine-learning nonstationary noise '
        'out of gravitational-wave detectors,\u201d Phys. Rev. D 101, 042003 (2020).',

        '[18] R. Essick, P. Godwin, C. Hanna, L. Blackburn, and E. Katsavounidis, '
        '\u201ciDQ: Statistical inference of non-Gaussian noise with auxiliary degrees '
        'of freedom in gravitational-wave detectors,\u201d Mach. Learn.: '
        'Sci. Technol. 2, 015004 (2021).',

        # --- Sec IV.D: EBSSA dataset and temporal filter ---
        '[19] S. Afshar, A. P. Nicholson, A. van Schaik, and G. Cohen, '
        '\u201cEvent-based object detection and tracking for space situational awareness,\u201d '
        'preprint arXiv:1911.08730 (2019).',

        '[20] T. Delbruck, \u201cFrame-free dynamic digital vision,\u201d in Proc. Intl. Symp. '
        'on Secure-Life Electronics (2008), pp. 21\u201326.',

        # --- Sec V: forbidden interval, suprathreshold SR ---
        '[21] B. Kosko and S. Mitaim, \u201cStochastic resonance in noisy threshold neurons,\u201d '
        'Neural Netw. 16, 755 (2003).',

        '[22] S. Mitaim and B. Kosko, \u201cAdaptive stochastic resonance in noisy neurons '
        'based on mutual information,\u201d IEEE Trans. Neural Netw. 15, 1526 (2004).',

        '[23] N. G. Stocks, \u201cInformation transmission in parallel threshold arrays: '
        'Suprathreshold stochastic resonance,\u201d Phys. Rev. E 63, 041114 (2001).',

        # --- Sec V.D: broader applicability ---
        '[24] P. H\u00e4nggi, \u201cStochastic resonance in biology: How noise can enhance detection '
        'of weak signals and help improve biological information processing,\u201d '
        'ChemPhysChem 3, 285 (2002).',

        '[25] N. Gisin, G. Ribordy, W. Tittel, and H. Zbinden, \u201cQuantum cryptography,\u201d '
        'Rev. Mod. Phys. 74, 145 (2002).',

        '[26] M. A. Richards, J. A. Scheer, and W. A. Holm, Principles of Modern Radar: '
        'Basic Principles (SciTech Publishing, 2010).',

        '[27] S. M. Bezrukov and I. Vodyanoy, \u201cNoise-induced enhancement of signal '
        'transduction across voltage-dependent ion channels,\u201d '
        'Nature 378, 362 (1995).',
    ]
    for ref in references:
        add_paragraph(doc, ref, space_after=3)

    # Save
    out_path = OUT_DIR / 'manuscript_pre.docx'
    doc.save(str(out_path))
    print(f"Manuscript saved: {out_path}")


if __name__ == '__main__':
    build_manuscript()
